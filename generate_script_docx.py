"""Generate presentation script Word document using python-docx."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "Spillover_Monitor_Presentation_Script.docx"

# ── colours ───────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0E, 0x11, 0x17)
GOLD   = RGBColor(0xC9, 0xA8, 0x4C)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0x1E, 0x23, 0x30)   # dark card bg for shading
MUTED  = RGBColor(0x88, 0x90, 0xA1)
BODY   = RGBColor(0x1A, 0x1F, 0x2E)   # near-black body text


# ── xml helpers ───────────────────────────────────────────────────────────────
def _rgb_hex(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def _set_para_shading(para, hex_color: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _para_spacing(para, before_pt=0, after_pt=6, line_pt=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(int(before_pt * 20)))
    spacing.set(qn("w:after"),  str(int(after_pt  * 20)))
    if line_pt:
        spacing.set(qn("w:line"),     str(int(line_pt * 20)))
        spacing.set(qn("w:lineRule"), "exact")
    pPr.append(spacing)


def _page_border(doc):
    """Add a gold top border to every page via sectPr."""
    sectPr = doc.sections[0]._sectPr
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")
    top = OxmlElement("w:top")
    top.set(qn("w:val"),   "single")
    top.set(qn("w:sz"),    "18")
    top.set(qn("w:space"), "24")
    top.set(qn("w:color"), _rgb_hex(GOLD))
    pgBorders.append(top)
    sectPr.append(pgBorders)


# ── paragraph builders ────────────────────────────────────────────────────────
def add_para(doc, text, bold=False, italic=False,
             color=None, size=10,
             align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=6,
             indent_cm=0):
    p = doc.add_paragraph()
    p.alignment = align
    _para_spacing(p, space_before, space_after)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_cover(doc):
    doc.add_paragraph()  # top padding
    doc.add_paragraph()

    # main title
    p = add_para(doc, "EQUITY-COMMODITIES SPILLOVER MONITOR",
                 bold=True, color=BODY, size=20,
                 align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=0, space_after=4)

    add_para(doc, "Presentation Script - Live Demo Walkthrough",
             color=MUTED, size=11,
             align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=0, space_after=4)

    add_para(doc, "Purdue University  ·  Daniels School of Business  ·  MGMT 69000-120",
             italic=True, color=GOLD, size=9,
             align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=0, space_after=6)

    # gold rule
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p, 2, 6)
    run = p.add_run("─" * 42)
    run.font.color.rgb = GOLD
    run.font.size = Pt(9)

    add_para(doc, "Heramb S. Patkar  ·  Jiahe Miao  ·  Ilian Zalomai",
             color=MUTED, size=10,
             align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=0, space_after=0)

    doc.add_page_break()


def add_section(doc, label, title, body_paras,
                action=None, note=None):
    # ── action cue (gold, italic, indented) ──────────────────────────────────
    if action:
        p = doc.add_paragraph()
        _para_spacing(p, before_pt=14, after_pt=2)
        p.paragraph_format.left_indent = Cm(0.4)
        run = p.add_run(f"→  {action}")
        run.italic = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = GOLD

    # ── label ─────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    _para_spacing(p, before_pt=2 if action else 16, after_pt=1)
    run = p.add_run(label.upper())
    run.font.size = Pt(7)
    run.font.color.rgb = GOLD
    run.font.bold = False
    # letter-spacing via rPr not easily supported; the all-caps + small size reads fine

    # ── title ─────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    _para_spacing(p, before_pt=1, after_pt=4)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = BODY

    # gold rule under title
    p = doc.add_paragraph()
    _para_spacing(p, 0, 6)
    run = p.add_run("─" * 72)
    run.font.color.rgb = GOLD
    run.font.size = Pt(7)

    # ── body paragraphs ───────────────────────────────────────────────────────
    for text in body_paras:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.right_indent = Cm(0.5)
        _para_spacing(p, before_pt=0, after_pt=7, line_pt=14)
        run = p.add_run(text)
        run.font.size = Pt(9.5)
        run.font.color.rgb = BODY

    # ── note ──────────────────────────────────────────────────────────────────
    if note:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        _para_spacing(p, before_pt=0, after_pt=4)
        run = p.add_run(f"▶  {note}")
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED


# ── content ──────────────────────────────────────────────────────────────────
SECTIONS = [
    {
        "label": "Opening - The Team",
        "title": "Who We Are",
        "body": [
            "I'm Heramb - equity research at Axis Securities, NISM-certified, BITS Pilani. "
            "Jiahe: capital markets and information systems from Kelley. "
            "Ilian: four years fintech risk ops and Deloitte banking consulting in Frankfurt. "
            "We built a cross-asset intelligence platform. Let me show you what it does.",
        ],
    },
    {
        "label": "Use Case 1 - Overview",
        "title": "\"What is happening in markets right now?\"",
        "action": "Navigate to: Overview  |  Start at the top of the page, scroll slowly downward",
        "body": [
            "This is your morning brief. Before you've touched anything: the current correlation "
            "regime, a live 0-to-100 risk score, best and worst performers across 15 equity "
            "indices and 17 commodity futures, and whether your equity-bond hedge is still intact.",

            "Scroll down - right here is a full briefing written by seven AI agents that ran "
            "before you opened the page. Macro context, geopolitical risk, commodities outlook, "
            "trade ideas, stress assessment. Two hours of analyst work. Every morning. "
            "On every load. Zero manual effort.",
        ],
        "note": "Point to: regime badge, risk score gauge, and the Risk Officer briefing block.",
    },
    {
        "label": "Use Case 2 - Portfolio Stress Test",
        "title": "\"My portfolio is down - have I seen this before?\"",
        "action": "Navigate to: Strategy → Portfolio Stress Test  |  Add 3-4 assets live",
        "body": [
            "Build your allocation here - any mix of equity indices, commodities, fixed income, "
            "or individual S&P 500 stocks. Hit run.",

            "The platform tests it against every major shock since 2008 - GFC, COVID, Ukraine, "
            "the Fed hiking cycle, the LME nickel squeeze. Pre-event, during, post: returns, "
            "max drawdown, Sharpe ratio per event. You immediately see whether today's move "
            "is a 2022 replay, or something you haven't encountered before.",
        ],
        "note": "Add S&P 500, Gold, TLT live. Point to the event heatmap after running.",
    },
    {
        "label": "Use Case 3 - War Impact Map",
        "title": "\"Which markets are exposed to what's happening geopolitically?\"",
        "action": "Navigate to: Analysis → War Impact Map  |  Switch to Combined view",
        "body": [
            "195 countries scored by equity market exposure to the three active conflict theaters "
            "right now - Ukraine, Gaza, and Iran-Hormuz. "
            "This column is the model's exposure score. This column is actual index performance "
            "since each conflict started. The gap between them - that's the mispricing.",

            "You're not just watching the news. You're seeing which markets the market "
            "hasn't caught up with yet.",
        ],
        "note": "Hover over a country in the top-25 table. Point to the exposure vs. repricing gap.",
    },
    {
        "label": "Use Case 4 - Scenario Engine",
        "title": "\"What does a Hormuz closure actually do to my book?\"",
        "action": "Navigate to: Strategy → Scenario Engine  |  Click preset: Geopolitical Escalation",
        "body": [
            "One click - oil shock, yield spike, credit spread widening, DXY move - "
            "propagated through live regression betas to every equity and commodity we cover. "
            "VaR and Expected Shortfall at 95 and 99 percent, with a waterfall showing "
            "which channel is driving which asset.",

            "You can customize every input or build a custom scenario from scratch. "
            "You walk into a risk committee with a number, not a gut feel.",
        ],
        "note": "Load the preset, let it run, point to the waterfall chart and VaR/ES table.",
    },
    {
        "label": "Use Case 5 - Trade Ideas",
        "title": "\"What should I actually do in this regime?\"",
        "action": "Navigate to: Strategy → Trade Ideas  |  Let the page load with live regime filter",
        "body": [
            "Trade cards are automatically filtered to today's live correlation regime - "
            "ideas that are structurally valid right now, not in a different market environment. "
            "Each card has a thesis, entry signal, risk consideration, and a live correlation "
            "chart confirming the pair relationship is intact today.",

            "The AI Trade Structurer synthesizes macro, geopolitical, stress, and commodities "
            "context simultaneously to generate new ideas on demand.",
        ],
        "note": "Point to the regime label on a trade card. Expand one card fully.",
    },
    {
        "label": "Use Case 6 - AI Analyst",
        "title": "\"I have a specific question. Get me an answer.\"",
        "action": "Navigate to: Research → AI Analyst  |  Type a live question",
        "body": [
            "Full dashboard state injected into every query - regime, risk score, all live "
            "returns, implied vol, active conflicts, and all seven agent outputs.",

            "Ask it: why is copper decoupling from equities right now? "
            "Or: does this setup look more like 2019 or 2022? "
            "Or: what does a Hormuz closure do to my Japan exposure? "
            "Answer in seconds, grounded in what's on this dashboard today.",
        ],
        "note": "Type a live question relevant to the current regime. Let it answer on screen.",
    },
    {
        "label": "There Is More",
        "title": "The Depth Is There - We Just Didn't Lead With It",
        "action": "Briefly scroll through: Correlation Analysis, Spillover Analytics, Strait Watch",
        "body": [
            "DCC-GARCH correlation regimes, Diebold-Yilmaz spillover networks, Markov "
            "regime transition forecasts, Granger causality grids, COT positioning overlays, "
            "geopolitical event forensics back to 2008, FRED macro intelligence, "
            "maritime chokepoint disruption scores - it is all here.",

            "If a customer gets interested, they can go as deep as they want. "
            "These pages don't need a tour. They reward exploration.",
        ],
    },
]

CLOSING = [
    "Finance has always had data. What it has never had is synthesis - "
    "fifteen equity markets, seventeen commodities, geopolitics, macro, and AI "
    "in one coherent view, updated continuously, with no manual work.",

    "That is what this platform does. "
    "If any of those use cases resonated, we are happy to go deeper. Questions?",
]


# ── build ─────────────────────────────────────────────────────────────────────
def build():
    doc = Document()

    # page setup - A4, 2.2 cm margins
    for section in doc.sections:
        section.page_width  = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.2)

    # default paragraph style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = BODY

    _page_border(doc)
    add_cover(doc)

    for sec in SECTIONS:
        add_section(
            doc,
            label=sec["label"],
            title=sec["title"],
            body_paras=sec["body"],
            action=sec.get("action"),
            note=sec.get("note"),
        )

    # closing
    p = doc.add_paragraph()
    _para_spacing(p, before_pt=18, after_pt=1)
    run = p.add_run("CLOSING")
    run.font.size = Pt(7)
    run.font.color.rgb = GOLD

    p = doc.add_paragraph()
    _para_spacing(p, before_pt=1, after_pt=4)
    run = p.add_run("The Pitch")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = BODY

    p = doc.add_paragraph()
    _para_spacing(p, 0, 6)
    run = p.add_run("─" * 72)
    run.font.color.rgb = GOLD
    run.font.size = Pt(7)

    for text in CLOSING:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.right_indent = Cm(0.5)
        _para_spacing(p, before_pt=0, after_pt=7, line_pt=14)
        run = p.add_run(text)
        run.font.size = Pt(9.5)
        run.font.color.rgb = BODY

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    build()
