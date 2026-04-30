"""
Generate Week 7 Capstone Final Review document (.docx) for industry panelists.
Reads CAPSTONE_FINAL.md and produces a formatted Word document.

Usage:
    python generate_capstone_docx.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

OUTPUT = "Spillover_Monitor_Capstone_Final.docx"

# ── Colours ───────────────────────────────────────────────────────────────────
GOLD  = RGBColor(0xCF, 0xB9, 0x91)   # Purdue gold
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BODY  = RGBColor(0x1A, 0x1A, 0x2E)   # near-black body text (readable on white)
MUTED = RGBColor(0x55, 0x5E, 0x70)   # muted grey
DARK  = RGBColor(0x0D, 0x0D, 0x1A)   # heading dark
ACCENT= RGBColor(0x8E, 0x6F, 0x3E)   # gold-brown accent

_GOLD_HEX  = "CFB991"
_DARK_HEX  = "0D0D1A"
_LIGHT_HEX = "F5F4F0"
_RULE_HEX  = "CFB991"


# ── XML helpers ───────────────────────────────────────────────────────────────
def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_bottom_border(paragraph, color_hex: str = _RULE_HEX, size: int = 6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_page_margins(doc, top=1.8, bottom=1.8, left=2.2, right=2.2):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)


# ── Style helpers ─────────────────────────────────────────────────────────────
def _run(para, text: str, bold=False, italic=False, size=11,
         color: RGBColor = None, font="Calibri"):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font
    if color:
        run.font.color.rgb = color
    return run


def _heading(doc, text: str, level=1):
    """Section heading styled with gold underline."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    para.paragraph_format.space_after  = Pt(4)
    if level == 1:
        _run(para, text, bold=True, size=15, color=DARK, font="Calibri")
        _add_bottom_border(para, _GOLD_HEX, 8)
    elif level == 2:
        _run(para, text, bold=True, size=12, color=DARK, font="Calibri")
        _add_bottom_border(para, "DDDDDD", 4)
    else:
        _run(para, text, bold=True, size=11, color=ACCENT, font="Calibri")
    return para


def _body(doc, text: str, size=10.5, indent_cm=0):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(4)
    if indent_cm:
        para.paragraph_format.left_indent = Cm(indent_cm)
    _run(para, text, size=size, color=BODY)
    return para


def _bullet(doc, text: str, level=0):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(2)
    para.paragraph_format.left_indent  = Cm(0.6 + level * 0.6)
    _run(para, text, size=10, color=BODY)
    return para


def _spacer(doc, size_pt=6):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(size_pt)
    para.paragraph_format.space_after  = Pt(0)
    return para


def _table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    # Header row
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        _set_cell_bg(cell, _DARK_HEX)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xCF, 0xB9, 0x91)
        run.font.name = "Calibri"

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        bg = _LIGHT_HEX if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            _set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = BODY
            run.font.name = "Calibri"

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)

    return tbl


# ── Cover page ────────────────────────────────────────────────────────────────
def _cover(doc):
    _spacer(doc, 40)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "Equity-Commodities Spillover Monitor", bold=True, size=22,
         color=DARK, font="Calibri")

    _spacer(doc, 6)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p2, "Week 7 Capstone — Final Industry Panelist Review", bold=False,
         size=14, color=ACCENT, font="Calibri")

    _spacer(doc, 4)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p3, "MGMT 69000-120: AI for Finance  ·  Purdue University Daniels School of Business",
         size=11, color=MUTED, font="Calibri")

    _spacer(doc, 16)
    # Gold rule
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_bottom_border(p4, _GOLD_HEX, 12)

    _spacer(doc, 20)
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p5, "Heramb S. Patkar  ·  Ilian Zalomai  ·  Jiahe Miao",
         bold=True, size=12, color=DARK, font="Calibri")

    _spacer(doc, 4)
    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p6, "Instructor: Prof. Cinder Zhang  ·  April 30, 2026  ·  300 Points",
         size=10, color=MUTED, font="Calibri")

    doc.add_page_break()


# ── Main document builder ─────────────────────────────────────────────────────
def build(doc: Document):
    _cover(doc)

    # ── Section 1: Executive Summary ─────────────────────────────────────────
    _heading(doc, "1. Executive Summary")
    _body(doc, (
        "The Equity-Commodities Spillover Monitor is an institutional-grade, multi-page analytical "
        "dashboard that answers a question standard financial terminals do not: how do geopolitical "
        "shocks transmit across asset classes in real time, and what does that mean for positioning?"
    ))
    _body(doc, (
        "The system ingests live data from six external sources, runs five independent econometric "
        "models, orchestrates seven AI agents in a dependency-ordered pipeline, and presents results "
        "across 22 pages in a purpose-built dark-theme Streamlit interface. Every analytical claim "
        "is backed by a published methodology; every model output is validated against a 56-case "
        "out-of-sample benchmark harness."
    ))
    _body(doc, (
        "Core value proposition: a single analyst can open this dashboard and, within two minutes, "
        "know what is driving cross-asset stress, which geopolitical conflicts are transmitting into "
        "commodity markets, and what trade structures are valid under the current correlation "
        "regime — grounded in real data and statistically rigorous methodology."
    ))

    # ── Section 2: Problem Statement ─────────────────────────────────────────
    _heading(doc, "2. Problem Statement & Market Gap")
    _heading(doc, "2.1 The Question This System Answers", level=2)
    _body(doc, (
        "How do geopolitical shocks transmit across asset classes — specifically between global "
        "equity markets and commodity futures — and what does that mean for portfolio positioning?"
    ))
    _body(doc, (
        "Standard market dashboards (Bloomberg, FactSet) answer: 'What are prices doing?' This "
        "system answers: 'Why are prices moving, where is the shock originating, and how does it "
        "propagate?'"
    ))

    _heading(doc, "2.2 Capabilities Comparison", level=2)
    _table(doc,
        ["Capability", "Bloomberg", "Standard Dashboard", "This System"],
        [
            ["Live price/return data", "✅", "✅", "✅"],
            ["Correlation levels", "✅", "✅", "✅"],
            ["Spillover directionality", "❌", "❌", "✅ Diebold-Yilmaz FEVD"],
            ["Regime-conditioned analysis", "Partial", "❌", "✅ 3-state adaptive"],
            ["Geopolitical → market attribution", "❌", "❌", "✅ CIS/TPS scoring"],
            ["AI-synthesized morning briefing", "❌", "❌", "✅ 7-agent pipeline"],
            ["Scenario propagation w/ regime conditioning", "❌", "❌", "✅ Nonlinear OLS betas"],
            ["Out-of-sample model validation", "❌", "❌", "✅ 56-case harness"],
        ],
        col_widths=[7.0, 2.5, 3.5, 5.0],
    )
    _spacer(doc)

    # ── Section 3: System Architecture ────────────────────────────────────────
    _heading(doc, "3. System Architecture")
    _heading(doc, "3.1 Three-Layer Design", level=2)
    _body(doc, (
        "The system is organized in three layers: Data Ingestion → Analytics Engine → AI Agent "
        "Workforce, with a Streamlit UI as the presentation layer."
    ))

    _heading(doc, "3.2 Data Pipeline — Six Live Sources", level=2)
    _table(doc,
        ["Source", "Data", "Cache TTL", "Notes"],
        [
            ["Yahoo Finance", "50+ tickers: equities, commodities, FX, vol indices", "900s", "Parallel-loaded via ThreadPoolExecutor"],
            ["FRED API", "24 macro series: yield curve, CPI, Fed Funds, GDP, PMI", "86400s", "Daily; last available for stale dates"],
            ["GDELT 2.0", "News volume and tone for 6 active conflicts", "7200s", "Conflict interest score = mentions × tone"],
            ["ACLED", "Georeferenced conflict events and fatalities", "21600s", "REST API; filtered to active conflict zones"],
            ["IMF PortWatch", "Vessel traffic at Hormuz + 5 chokepoints", "86400s", "AIS-derived; deviation from 90d baseline"],
            ["EIA", "Weekly US petroleum inventory", "86400s", "Wednesday release; surprise vs consensus"],
        ],
        col_widths=[2.8, 5.2, 2.0, 5.0],
    )
    _spacer(doc)

    _heading(doc, "3.3 Analytics Engine — Five Independent Models", level=2)

    _heading(doc, "A. Diebold-Yilmaz FEVD Spillover Index", level=3)
    _body(doc, (
        "The headline methodology. VAR(p) with BIC-optimal lag selection over a 252-day rolling "
        "window. 10-step generalized forecast-error variance decomposition (GFEVD) produces: "
        "pairwise spillover matrix, net directional spillover, and Total Spillover Index (TSI). "
        "Key design choice: generalized (not Cholesky) decomposition eliminates ordering dependence."
    ))

    _heading(doc, "B. Granger Causality — BIC Lag + Holm-Bonferroni Correction", level=3)
    _body(doc, (
        "VAR(p) fitted with BIC for lag selection before testing. Holm-Bonferroni step-down "
        "correction applied across the full N×M×2 test grid to control family-wise error rate. "
        "Results visualized as a directed causality matrix with net directional arrows."
    ))

    _heading(doc, "C. Transfer Entropy (Schreiber, 2000)", level=3)
    _body(doc, (
        "Non-parametric directed information measure capturing nonlinear dependencies that Granger "
        "misses. Lag-optimized (search over lags 1–7). 200-permutation shuffle test establishes "
        "significance threshold; only significant edges retained."
    ))

    _heading(doc, "D. DCC-GARCH Dynamic Correlation (Engle, 2002)", level=3)
    _body(doc, (
        "Two-step estimation: EWMA pre-whitening (λ=0.94, RiskMetrics) removes heteroskedasticity "
        "contamination before DCC(1,1) recursion. Raw-return DCC produces biased correlations "
        "during high-vol periods — the pre-whitening step is the critical design choice."
    ))

    _heading(doc, "E. Composite Risk Score — GRS (0–100)", level=3)
    _body(doc, "Three-layer architecture:")
    _bullet(doc, "GRS = 40% × CIS (Conflict Intensity) + 35% × TPS (Transmission Pressure) + 25% × MCS (Market Conditions)")
    _bullet(doc, "HHI-based breadth multiplier (n_eff^0.25) prevents single-conflict concentration from inflating composite")
    _bullet(doc, "3-state correlation regime: Decoupled (ρ<0.30) / Transitioning (0.30–0.55) / High Coupling (ρ≥0.55)")
    _spacer(doc)

    _heading(doc, "3.4 AI Agent Workforce — 7 Agents, 4-Round Pipeline", level=2)
    _table(doc,
        ["Round", "Agent", "Role", "Dependencies"],
        [
            ["1 (parallel)", "macro_strategist", "Macro regime, yield curve, Fed outlook", "None"],
            ["1 (parallel)", "geopolitical_analyst", "Active conflict summary, CIS/TPS attribution", "None"],
            ["1 (parallel)", "commodities_specialist", "Supply/demand dynamics per commodity", "None"],
            ["2 (sequential)", "risk_officer", "Synthesizes Round 1 → composite risk narrative + 3 risk flags", "Round 1 all"],
            ["3 (parallel)", "stress_engineer", "Worst-case scenario, tail risk quantification", "risk_officer"],
            ["3 (parallel)", "signal_auditor", "Data quality flags, stale data warnings, model limitations", "risk_officer"],
            ["4 (sequential)", "trade_structurer", "3 regime-conditioned trade ideas with entry/exit/rationale", "Round 3 all"],
            ["4 (sequential)", "quality_officer (CQO)", "Schema validation; routes failures to remediation loop", "Round 3 all"],
        ],
        col_widths=[2.5, 3.0, 5.5, 3.0],
    )
    _spacer(doc)
    _body(doc, (
        "All agents receive the full analytics context (GRS, CIS, TPS, regime, DY spillover TSI) "
        "as structured input — not raw text. Pydantic-validated outputs enforce schema at every "
        "agent boundary. The CQO remediation loop automatically routes quality flags to the "
        "relevant specialist for active correction (up to 2 rounds)."
    ))

    # ── Section 4: Dashboard ──────────────────────────────────────────────────
    _heading(doc, "4. Dashboard — 22 Pages")
    _table(doc,
        ["Page", "Purpose", "Key Analytics"],
        [
            ["Command Center", "Home: live intelligence, GRS, conflict landscape, risk compass", "GRS, CIS, TPS, regime, 5-axis radar, returns heatmap, transmission channels"],
            ["Correlation Analysis", "Rolling equity↔commodity correlation; regime time series", "DCC-GARCH, Pearson rolling, 3-state regime classifier"],
            ["Spillover Analytics", "Diebold-Yilmaz network; pairwise FEVD; rolling TSI", "D-Y (2012), VAR BIC-optimal, directional arrows"],
            ["Macro Intelligence", "Yield curve, CPI, Fed Funds, PMI, GDP", "FRED live data, yield curve spread, macro regime"],
            ["Geopolitical Triggers", "Per-conflict CIS/TPS timelines; GDELT/ACLED overlays", "Conflict Intelligence Score, lead-lag signal"],
            ["War Impact Map", "Country-level commodity exposure to active conflicts", "War scoring framework (Zalomai), concurrent-war amplifier"],
            ["Strait Watch", "Maritime chokepoint disruption; Hormuz crisis tracker", "PortWatch AIS, EIA petroleum inventory"],
            ["Scenario Engine", "Custom shock propagation; 13 historical scenarios", "Parametric shock model, regime-nonlinear OLS betas"],
            ["Trade Ideas", "6 regime-conditioned trade structures", "Correlation regime + AI trade structurer"],
            ["Portfolio Stress Test", "CSV upload; portfolio-weighted impact attribution", "Exposure scoring (SES/TAE/SAS)"],
            ["Watchlist", "17-commodity futures; backwardation/contango curve", "Futures curve analysis"],
            ["Transmission Matrix", "Full pairwise Granger matrix; network view", "Granger + Holm-Bonferroni, directional arrows"],
            ["Exposure Scoring", "SES/TAE/SAS per asset-conflict pair", "Structural + transmission-adjusted exposure"],
            ["Conflict Intelligence", "GDELT + ACLED live monitor; news feed", "GDELT 2.0 REST, conflict timeline"],
            ["Threat Activity Monitor", "Real-time ACLED threat event stream", "ACLED REST API"],
            ["Model Accuracy", "56-case benchmark harness results per agent", "Out-of-sample validation, documented pass/fail"],
            ["AI Chat", "Natural language interface to all analytics", "Claude Sonnet / GPT-4o"],
            ["About — AI Workforce", "7-agent documentation and pipeline diagram", "—"],
            ["Methodology", "Full statistical methodology with citations", "All methods documented"],
        ],
        col_widths=[3.5, 5.5, 6.0],
    )
    _spacer(doc)
    doc.add_page_break()

    # ── Section 5: Benchmark Validation ───────────────────────────────────────
    _heading(doc, "5. Benchmark Validation — Model Accuracy")
    _body(doc, (
        "The 56-case out-of-sample benchmark harness (evals/run_eval.py) validates each agent's "
        "outputs against historical ground truth across 15 well-documented market events from "
        "2008 to 2024. Pass threshold: 60% hit rate, set at project outset."
    ))

    _heading(doc, "5.1 Results (2026-04-28)", level=2)
    _table(doc,
        ["Agent", "Cases", "Passed", "Hit Rate", "60% Gate"],
        [
            ["risk_officer", "30", "13", "43.3%", "❌"],
            ["macro_strategist", "11", "7", "63.6%", "✅"],
            ["geopolitical_analyst", "11", "0", "0.0%*", "❌"],
            ["commodities_specialist", "3", "1", "33.3%", "❌"],
            ["signal_auditor", "4", "0", "0.0%", "❌"],
            ["stress_engineer", "1", "1", "100.0%", "✅"],
        ],
        col_widths=[4.5, 2.0, 2.0, 2.5, 2.5],
    )
    _spacer(doc)

    _heading(doc, "5.2 Interpretation", level=2)
    _bullet(doc, (
        "macro_strategist (63.6%) and stress_engineer (100%) pass the gate. The macro model "
        "correctly classifies regimes and yield curve states across COVID, SVB, and other "
        "macro-driven events."
    ))
    _bullet(doc, (
        "risk_officer (43.3%) — primary failure mode is over-scoring calm periods. The composite "
        "GRS is geopolitically anchored; when financial markets stabilize but conflicts continue, "
        "GRS remains elevated. Acknowledged calibration issue, not a data error."
    ))
    _bullet(doc, (
        "geopolitical_analyst (0%*) — all failures are 'None' actual values. This is a data "
        "infrastructure issue: score_all_conflicts() requires live GDELT access, which cannot "
        "be replayed historically. CIS/TPS cannot be computed for past dates. Documented as "
        "known limitation in evals/results-2026-04.md."
    ))
    _bullet(doc, (
        "Lehman Collapse (2008) is unscorable — FRED macro data unavailable before 2010 in the "
        "required format. GRS returns nan; documented in benchmark results."
    ))

    # ── Section 6: Performance ────────────────────────────────────────────────
    _heading(doc, "6. Performance Architecture")
    _heading(doc, "6.1 Cold-Start Optimization", level=2)
    _body(doc, "Before optimization: ~30s sequential cold start. After optimization: ~15s via parallelization.")
    _body(doc, (
        "ThreadPoolExecutor(max_workers=2) parallelizes score_all_conflicts() (GDELT HTTP, ~10s) "
        "and _load_market_pulse() (yfinance, ~8s) while _load_market_risk() runs on the main "
        "thread. Reason: _load_market_risk() writes to st.session_state — unsafe to thread in "
        "Streamlit."
    ))
    _heading(doc, "6.2 Cache TTL Configuration", level=2)
    _table(doc,
        ["Function", "Cache TTL", "Rationale"],
        [
            ["score_all_conflicts", "7200s (2h)", "GDELT updates at 15-min intervals; 2h avoids API throttling"],
            ["_load_market_risk", "900s (15min)", "Matches intraday signal relevance"],
            ["_load_market_pulse", "900s (15min)", "Liquid tickers; 15min lag acceptable"],
            ["FRED series", "86400s (24h)", "Daily release; no intraday updates"],
            ["ACLED", "21600s (6h)", "ACLED bulk data updates daily"],
        ],
        col_widths=[4.0, 2.8, 8.2],
    )
    _spacer(doc)
    _body(doc, "Warm-cache load (all subsequent navigations): <2s for all pages.")

    # ── Section 7: Business Value ─────────────────────────────────────────────
    _heading(doc, "7. Business Value & Impact")

    _heading(doc, "7.1 For Institutional Analysts", level=2)
    _bullet(doc, (
        "Time savings: A cross-asset risk briefing that takes 60–90 minutes to assemble manually "
        "(Bloomberg + Excel + manual GDELT) is produced in under 15 seconds on dashboard load."
    ))
    _bullet(doc, (
        "Signal quality: The AI pipeline synthesizes five independent model outputs into one "
        "coherent narrative, with explicit quality flagging by the CQO."
    ))
    _bullet(doc, (
        "Geopolitical attribution: Standard terminals cannot answer 'how much of today's WTI "
        "move is the Red Sea vs. the Fed?' This system quantifies that via CIS × TPS channel "
        "weighting."
    ))

    _heading(doc, "7.2 For Portfolio Managers", level=2)
    _bullet(doc, (
        "Regime-conditioned trade ideas: 6 trade structures explicitly conditioned on the "
        "current correlation regime. A trade valid in Regime 1 (Decoupled) is explicitly "
        "invalidated in Regime 3 (High Coupling)."
    ))
    _bullet(doc, (
        "Scenario propagation: Custom shocks propagated through regime-conditioned OLS betas. "
        "'Hormuz blockade — 30% oil supply cut' → portfolio P&L impact in seconds."
    ))

    _heading(doc, "7.3 For Risk Officers", level=2)
    _bullet(doc, (
        "GRS: A single 0–100 number with documented decomposition (CIS/TPS/MCS) and full "
        "methodology transparency. All inputs reproducible from public data sources."
    ))
    _bullet(doc, (
        "Out-of-sample validation: 56-case harness gives quantitative confidence bounds on "
        "model accuracy — documented, not claimed."
    ))

    _heading(doc, "7.4 Scalability Path", level=2)
    _bullet(doc, "Additional conflicts: add entry to CONFLICT_REGISTRY in config.py — no code changes")
    _bullet(doc, "Additional data sources: add @st.cache_data-decorated loader in src/data/loader.py")
    _bullet(doc, "Additional AI agents: register in agent_orchestrator.py with dependency declaration")
    _bullet(doc, "Production deployment: render.yaml configures Render.com auto-deploy on git push")

    # ── Section 8: Demo Script ────────────────────────────────────────────────
    _heading(doc, "8. End-to-End Demo Script")
    _body(doc, "Estimated duration: 12–15 minutes for full panel review.")

    _heading(doc, "Phase 1: Command Center (3 min)", level=2)
    _bullet(doc, "GRS gauge: 'This is the composite risk score — 40% geopolitical intensity, 35% transmission pressure, 25% market conditions.'")
    _bullet(doc, "Conflict Landscape: 'Each dot is a tracked conflict. Top-right quadrant = critical (high intensity, high transmission).'")
    _bullet(doc, "Risk Compass: 'Five-axis radar — CIS, TPS, MCS, VIX-normalized Volatility, Coupling. Large red polygon = multi-channel stress.'")
    _bullet(doc, "Transmission Channels: 'CIS-weighted aggregate showing which channels carry most stress: energy, shipping, FX, equity/inflation.'")

    _heading(doc, "Phase 2: Spillover Analytics (3 min)", level=2)
    _bullet(doc, "Diebold-Yilmaz matrix: 'Read row-by-row: X% of WTI forecast variance explained by shocks to S&P 500.'")
    _bullet(doc, "Rolling TSI: 'Total Spillover Index — integration spikes during COVID and Ukraine when normally-decoupled markets synchronize.'")
    _bullet(doc, "Net directional arrows: 'Right now, net flow is commodity→equity (supply shock) or equity→commodity (demand destruction).'")

    _heading(doc, "Phase 3: Geopolitical Layer (2 min)", level=2)
    _bullet(doc, "War Impact Map: 'Concurrent-war amplifier — Ukraine cutting wheat while Red Sea disrupts shipping is multiplicative in some commodities.'")
    _bullet(doc, "Strait Watch: 'Live AIS vessel traffic at Hormuz and five chokepoints. Disruption = deviation from 90-day baseline.'")

    _heading(doc, "Phase 4: Scenario Engine (2 min)", level=2)
    _bullet(doc, "Select 'Hormuz Blockade': 'In Regime 3 (High Coupling), the equity impact is 1.4× larger than in Regime 1.'")
    _bullet(doc, "Compound scenario: 'Stack Hormuz blockade + Fed rate cut. System handles nonlinear interaction via regime conditioning.'")

    _heading(doc, "Phase 5: AI Agent Output (2 min)", level=2)
    _bullet(doc, "Morning briefing: 'Seven agents, four rounds, one coherent narrative — Risk Officer synthesized macro + geopolitical + commodity inputs.'")
    _bullet(doc, "Model Accuracy: '56 cases, 15 historical events. Macro_strategist passes at 63.6%. Geopolitical agent fails — documented infrastructure limitation, not a hidden bug.'")

    # ── Section 9: Methodological Integrity ───────────────────────────────────
    _heading(doc, "9. Methodological Integrity")
    _heading(doc, "9.1 All 25 AUDIT.md Gaps Resolved", level=2)
    _body(doc, (
        "The initial gap analysis identified 25 deficiencies against the core problem statement. "
        "All 25 are marked ✅ Fixed in AUDIT.md. Key resolutions:"
    ))
    _bullet(doc, "Gap 15 (circular ground truth): historical benchmark rewritten from non-VIX sources")
    _bullet(doc, "Gap 18 (no out-of-sample test): 56-case harness with documented pass/fail criteria")
    _bullet(doc, "Gap 20 (TE lag=1 hardcoded): lag search over 1–7, commodity→equity lag 2–5 days")
    _bullet(doc, "Gap 19 (linear betas): regime-conditioned OLS betas replace linear scenario propagation")
    _bullet(doc, "Gap 24 (Granger not network-visualized): Transmission Matrix page with directional arrows")
    _bullet(doc, "Gap 16 (silent failures): stale data banners and insufficient data warnings throughout")

    _heading(doc, "9.2 Known Limitations — Documented, Not Hidden", level=2)
    _bullet(doc, "risk_officer over-scores calm periods: GRS geopolitically anchored; fix requires separate geo/market composite")
    _bullet(doc, "geopolitical_analyst requires live data: CIS/TPS not replayable historically (GDELT is real-time)")
    _bullet(doc, "Lehman 2008 unscorable: FRED data unavailable pre-2010 in required format")
    _bullet(doc, "Transfer entropy at 200 shuffles: 1000-shuffle test would increase power; deferred due to runtime cost")

    # ── Section 10: References ────────────────────────────────────────────────
    _heading(doc, "10. Academic References")
    refs = [
        ("Diebold & Yilmaz (2012)", "International Journal of Forecasting", "Better to give than to receive: Predictive directional measurement of volatility spillovers"),
        ("Granger (1969)", "Econometrica", "Investigating causal relations by econometric models and cross-spectral methods"),
        ("Lütkepohl (2005)", "Springer: New Introduction to Multiple Time Series Analysis", "VAR lag selection — Chapter 4"),
        ("Schreiber (2000)", "Physical Review Letters", "Measuring information transfer (Transfer Entropy)"),
        ("Engle (2002)", "Journal of Business & Economic Statistics", "Dynamic conditional correlation: A simple class of multivariate GARCH models"),
        ("Hamilton (1989)", "Econometrica", "A new approach to the economic analysis of nonstationary time series"),
        ("Illing & Liu (2006)", "Journal of Financial Stability", "Measuring financial stress in a developed country: An application to Canada"),
        ("J.P. Morgan / Reuters (1996)", "RiskMetrics — Technical Document, 4th ed.", "EWMA volatility (RiskMetrics)"),
        ("Holm (1979)", "Scandinavica Statistica", "A simple sequentially rejective multiple test procedure"),
        ("Hamilton (1994)", "Princeton UP: Time Series Analysis", "Markov regime-switching dynamics — Chapter 22"),
    ]
    _table(doc,
        ["Author & Year", "Source", "Title / Contribution"],
        refs,
        col_widths=[4.0, 4.5, 6.5],
    )
    _spacer(doc)

    # ── Section 11: Team ─────────────────────────────────────────────────────
    _heading(doc, "11. Team Contributions")
    _table(doc,
        ["Contributor", "Core Deliverables"],
        [
            ["Heramb S. Patkar", (
                "Full system architecture and build; all 22 pages; AI agent orchestrator (7 agents, CQO loop); "
                "five analytics modules (D-Y, Granger, TE, DCC-GARCH, GRS); all data integrations "
                "(yfinance, FRED, GDELT, ACLED, PortWatch, EIA); 56-case benchmark harness; "
                "deployment pipeline (Render.com); all submission documents"
            )],
            ["Ilian Zalomai", (
                "War Impact Map scoring framework; 13 geopolitical event catalog; Strait Watch chokepoint "
                "framework; Iran/Hormuz crisis content; concurrent-war amplifier methodology; "
                "all dashboard narrative text connecting quantitative outputs to market implications"
            )],
            ["Jiahe Miao", (
                "Correlation regime taxonomy (4-state); rolling window methodology; 6 regime-conditioned "
                "trade structures; Fixed Income signal framework; D-Y methodology research; "
                "private credit proxy selection; equity/commodity return data quality review"
            )],
        ],
        col_widths=[4.0, 11.0],
    )
    _spacer(doc)

    # ── Footer info ───────────────────────────────────────────────────────────
    _spacer(doc, 12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_bottom_border(p, _GOLD_HEX, 4)

    _spacer(doc, 4)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p2,
         "GitHub: github.com/HPATKAR/equity-commodities-spillover  ·  "
         "Live: equity-commodities-spillover.onrender.com  ·  2026-04-30",
         size=9, color=MUTED, font="Calibri")


# ── Entry point ────────────────────────────────────────────────────────────────
def main():
    doc = Document()
    _set_page_margins(doc)

    # Default body font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    build(doc)
    doc.save(OUTPUT)
    print(f"Saved → {OUTPUT}")


if __name__ == "__main__":
    main()
