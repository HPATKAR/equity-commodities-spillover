"""
Generate formal submission Word document (.docx):
  Part I  - Research Brief
  Part II - User Manual
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT = "Spillover_Monitor_Submission.docx"

# ── colours ──────────────────────────────────────────────────────────────────
# Dark dashboard theme: page bg #0E1117, cards #1E2330, gold accents
NAVY   = RGBColor(0xFF, 0xFF, 0xFF)   # section/part headings → white on dark bg
GOLD   = RGBColor(0xC9, 0xA8, 0x4C)  # gold accent (unchanged)
BODY   = RGBColor(0xD4, 0xD8, 0xE2)  # body text → light grey-white
MUTED  = RGBColor(0x88, 0x90, 0xA1)  # subdued text → mid grey
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)  # pure white (unchanged)
LIGHT  = RGBColor(0x1E, 0x23, 0x30)  # card / callout bg → dark blue-grey
RULE_C = RGBColor(0x2A, 0x30, 0x45)  # subtle rule → dark slate

# Hex constants used directly in XML calls
_PAGE_BG   = "0E1117"   # page background
_CARD_BG   = "1E2330"   # callout / table-header background
_ZEBRA_BG  = "161B27"   # alternate table row
_BORDER_DK = "2A3045"   # cell border (subtle)
_BORDER_MD = "1E2330"   # cell left/right border


# ── XML helpers ───────────────────────────────────────────────────────────────
def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, **kwargs):
    """Add borders to a cell. kwargs: top/bottom/left/right = (size, color_hex)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, (sz, color) in kwargs.items():
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _keep_with_next(para):
    pPr = para._p.get_or_add_pPr()
    kwn = OxmlElement("w:keepNext")
    pPr.append(kwn)


def _keep_together(para):
    pPr = para._p.get_or_add_pPr()
    kl = OxmlElement("w:keepLines")
    pPr.append(kl)


def _page_break_before(para):
    pPr = para._p.get_or_add_pPr()
    pb = OxmlElement("w:pageBreakBefore")
    pPr.append(pb)


def _keep_table_together(table):
    """Mark every row to not break across pages."""
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit")
        trPr.append(cant)


def _repeat_header(table):
    """Repeat the first row as header on each page."""
    trPr = table.rows[0]._tr.get_or_add_trPr()
    tblH = OxmlElement("w:tblHeader")
    trPr.append(tblH)


def _add_bottom_border(para, color="C9A84C", size=12):
    """Gold underline via paragraph bottom border."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(size))
    bot.set(qn("w:color"), color)
    bot.set(qn("w:space"), "4")
    pBdr.append(bot)
    pPr.append(pBdr)


def _set_para_spacing(para, before=0, after=6, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def _shade_row(row, hex_color):
    for cell in row.cells:
        _set_cell_bg(cell, hex_color)


def _set_page_background(doc, hex_color=_PAGE_BG):
    """Fill the entire page with a background colour (requires displayBackgroundShape in settings)."""
    bg = OxmlElement("w:background")
    bg.set(qn("w:color"), hex_color)
    bg.set(qn("w:themeColor"), "none")
    # w:background must precede w:body in the document element
    doc.element.insert(0, bg)
    # Enable rendering in settings.xml
    settings_elem = doc.settings.element
    disp = OxmlElement("w:displayBackgroundShape")
    settings_elem.insert(0, disp)


# ── document setup ────────────────────────────────────────────────────────────
def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)
    sec.top_margin    = Cm(2.2)
    sec.bottom_margin = Cm(2.2)

    # default normal style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.font.color.rgb = BODY
    normal.paragraph_format.space_after = Pt(6)

    _set_page_background(doc)
    return doc


# ── paragraph helpers ─────────────────────────────────────────────────────────
def add_para(doc, text, bold=False, italic=False, color=None,
             size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=6, keep_next=False):
    p = doc.add_paragraph()
    p.alignment = align
    _set_para_spacing(p, before=int(space_before * 20),
                      after=int(space_after * 20))
    if keep_next:
        _keep_with_next(p)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    return p


def add_mixed_para(doc, parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                   space_before=0, space_after=6):
    """parts = list of (text, bold, italic, color, size)"""
    p = doc.add_paragraph()
    p.alignment = align
    _set_para_spacing(p, before=int(space_before * 20),
                      after=int(space_after * 20))
    for text, bold, italic, color, size in parts:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = "Calibri"
        if color:
            run.font.color.rgb = color
    return p


def add_bullet(doc, text, level=0, space_after=3):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_para_spacing(p, after=int(space_after * 20))
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.name = "Calibri"
    run.font.color.rgb = BODY
    return p


def add_space(doc, pt=6):
    p = doc.add_paragraph()
    _set_para_spacing(p, before=0, after=0)
    run = p.add_run()
    run.font.size = Pt(pt)
    return p


# ── heading helpers ───────────────────────────────────────────────────────────
def add_part_header(doc, part_label, title):
    doc.add_page_break()
    p_label = add_para(doc, part_label.upper(), bold=False, color=GOLD,
                       size=8, align=WD_ALIGN_PARAGRAPH.LEFT,
                       space_before=0, space_after=2)
    p_title = add_para(doc, title, bold=True, color=NAVY,
                       size=18, align=WD_ALIGN_PARAGRAPH.LEFT,
                       space_before=0, space_after=4)
    _add_bottom_border(p_title, color="C9A84C", size=18)
    _add_bottom_border(p_title, color="C9A84C", size=6)
    add_para(doc, "", size=9)


def add_section_heading(doc, number, title, subtitle=None):
    p_label = add_para(doc, f"SECTION {number}", bold=False, color=GOLD,
                       size=7.5, align=WD_ALIGN_PARAGRAPH.LEFT,
                       space_before=12, space_after=1, keep_next=True)
    p_title = add_para(doc, title, bold=True, color=NAVY,
                       size=14, align=WD_ALIGN_PARAGRAPH.LEFT,
                       space_before=0, space_after=3, keep_next=True)
    _add_bottom_border(p_title, color="C9A84C", size=10)
    if subtitle:
        add_para(doc, subtitle, italic=True, color=MUTED,
                 size=9, space_before=2, space_after=8)
    add_space(doc, 4)


def add_sub_heading(doc, title):
    p = add_para(doc, title, bold=True, color=GOLD,
                 size=11, align=WD_ALIGN_PARAGRAPH.LEFT,
                 space_before=10, space_after=2, keep_next=True)
    _add_bottom_border(p, color="2A3045", size=6)
    return p


def add_subsub_heading(doc, title):
    p = add_para(doc, title, bold=True, color=WHITE,
                 size=10, align=WD_ALIGN_PARAGRAPH.LEFT,
                 space_before=8, space_after=2, keep_next=True)
    return p


def add_table_caption(doc, text):
    add_para(doc, text, bold=True, color=MUTED,
             size=9, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=6, space_after=3, keep_next=True)


def add_callout(doc, label, text, bg_hex=_CARD_BG, border_hex="C9A84C"):
    """Shaded callout box via a single-cell table."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, bg_hex)
    _set_cell_border(cell,
        top=(6, border_hex), bottom=(6, border_hex),
        left=(24, border_hex), right=(6, _CARD_BG))
    cell.width = Inches(6.1)

    # label line
    p_lbl = cell.add_paragraph()
    _set_para_spacing(p_lbl, before=80, after=40)
    r = p_lbl.add_run(label)
    r.bold = True
    r.font.size = Pt(8)
    r.font.name = "Calibri"
    r.font.color.rgb = GOLD

    # body line
    p_body = cell.add_paragraph()
    _set_para_spacing(p_body, before=0, after=80)
    r2 = p_body.add_run(text)
    r2.font.size = Pt(9.5)
    r2.font.name = "Calibri"
    r2.font.color.rgb = BODY
    p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # remove default empty first para
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
    add_space(doc, 6)


# ── table helper ──────────────────────────────────────────────────────────────
def add_styled_table(doc, data, col_widths_cm, zebra=True):
    """
    data[0] = header row (list of strings).
    col_widths_cm = list of column widths in cm.
    Text auto-wraps. Repeats header on each page.
    """
    n_cols = len(data[0])
    tbl = doc.add_table(rows=len(data), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"
    _repeat_header(tbl)

    for r_idx, row_data in enumerate(data):
        row = tbl.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            # set width
            cell.width = Cm(col_widths_cm[c_idx])

            # style the paragraph inside the cell
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_para_spacing(p, before=60, after=60)
            run = p.add_run(str(cell_text))
            run.font.name = "Calibri"

            if r_idx == 0:
                # header
                _set_cell_bg(cell, _CARD_BG)
                _set_cell_border(cell,
                    top=(6, _CARD_BG), bottom=(12, "C9A84C"),
                    left=(6, _CARD_BG), right=(6, _CARD_BG))
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = WHITE
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            else:
                run.font.size = Pt(9)
                run.font.color.rgb = BODY
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                # zebra
                if zebra and r_idx % 2 == 0:
                    _set_cell_bg(cell, _ZEBRA_BG)
                # bottom rule
                _set_cell_border(cell,
                    bottom=(4, _BORDER_DK),
                    left=(4, _BORDER_MD), right=(4, _BORDER_MD))

    add_space(doc, 6)
    return tbl


# ════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════
def build_cover(doc):
    for _ in range(8):
        add_space(doc, 12)

    # gold accent line
    p_accent = doc.add_paragraph()
    _add_bottom_border(p_accent, color="C9A84C", size=18)
    _set_para_spacing(p_accent, before=0, after=120)

    add_para(doc, "EQUITY-COMMODITIES SPILLOVER MONITOR",
             bold=True, color=WHITE, size=22,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "Research Brief & User Manual",
             bold=False, color=GOLD, size=13,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)

    p_rule2 = doc.add_paragraph()
    _add_bottom_border(p_rule2, color="C9A84C", size=6)
    _set_para_spacing(p_rule2, before=60, after=120)

    add_para(doc, "Course Project 3  ·  MGMT 69000-120: AI for Finance",
             color=MUTED, size=10,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    add_para(doc, "Purdue University  ·  Mitchell E. Daniels, Jr. School of Business",
             color=MUTED, size=10,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

    for _ in range(4):
        add_space(doc, 12)

    add_para(doc, "Heramb S. Patkar  ·  Jiahe Miao  ·  Ilian Zalomai",
             bold=True, color=WHITE, size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    add_para(doc, "April 2026",
             color=MUTED, size=10,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════════════
def build_toc(doc):
    add_para(doc, "TABLE OF CONTENTS", bold=False, color=GOLD, size=8,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=2)
    p = add_para(doc, "Contents", bold=True, color=WHITE, size=14,
                 space_before=0, space_after=3)
    _add_bottom_border(p, color="C9A84C", size=10)
    add_space(doc, 6)

    toc_entries = [
        (True, "PART I - RESEARCH BRIEF", 0),
        (False, "1.  Landscape Research", 1),
        (False, "    1.1  Financial Analytics Industry Overview", 2),
        (False, "    1.2  Cross-Asset Spillover Methods: Academic State of the Art", 2),
        (False, "    1.3  AI in Financial Analytics", 2),
        (False, "    1.4  Competitive Landscape", 2),
        (False, "2.  Gap Analysis", 1),
        (False, "    2.1  Identified Gaps in Existing Solutions", 2),
        (False, "    2.2  Opportunity Summary", 2),
        (False, "3.  Project Definition and Scope", 1),
        (False, "    3.1  Problem Statement", 2),
        (False, "    3.2  Objectives", 2),
        (False, "    3.3  Coverage Universe", 2),
        (False, "    3.4  Scope Boundaries", 2),
        (False, "4.  Technology Choices and Justification", 1),
        (False, "    4.1  Technology Stack Overview", 2),
        (False, "    4.2  Justification by Layer", 2),
        (True, "PART II - USER MANUAL", 0),
        (False, "5.  Getting Started", 1),
        (False, "    5.1  System Requirements & Installation", 2),
        (False, "    5.2  Configuration (API Keys)", 2),
        (False, "    5.3  Navigation Overview", 2),
        (False, "6.  Page-by-Page Operating Guide", 1),
        (False, "7.  AI Workforce Reference", 1),
        (False, "8.  Interpreting Key Outputs", 1),
        (False, "9.  Troubleshooting & FAQ", 1),
    ]

    for is_part, text, level in toc_entries:
        if is_part:
            add_space(doc, 4)
            add_para(doc, text, bold=True, color=GOLD, size=10,
                     space_before=2, space_after=2)
        elif level == 1:
            add_para(doc, text, bold=True, color=WHITE, size=9.5,
                     space_before=1, space_after=1)
        else:
            add_para(doc, text, color=MUTED, size=9,
                     space_before=0, space_after=1)

    doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# PART I
# ════════════════════════════════════════════════════════════════════════════
def build_part1_intro(doc):
    add_part_header(doc, "Part I", "Research Brief")
    add_para(doc,
        "This research brief documents the intellectual foundation, design rationale, and "
        "technology architecture of the Equity-Commodities Spillover Monitor. It is structured "
        "to satisfy the four required deliverables: landscape research, gap analysis, project "
        "definition and scope, and technology choices with justification.",
        size=10, space_after=8)


def build_section1(doc):
    add_section_heading(doc, "1", "Landscape Research",
        "Current state of the financial analytics industry, academic spillover methodology, "
        "and AI integration in finance.")

    add_sub_heading(doc, "1.1  Financial Analytics Industry Overview")
    add_para(doc,
        "The professional financial analytics market is currently dominated by a small number "
        "of high-cost terminal-based platforms. Bloomberg Terminal, with approximately 330,000 "
        "active subscriptions at roughly USD 25,000 per seat per annum, commands the largest "
        "institutional footprint. Refinitiv Eikon (now LSEG Workspace), FactSet Research "
        "Systems, and S&P Global Market Intelligence collectively serve the remainder of the "
        "institutional buy-side and sell-side. These platforms are characterised by "
        "comprehensive data access but are cost-prohibitive for academic institutions, "
        "emerging market practitioners, and smaller investment firms.")
    add_para(doc,
        "Within the quantitative finance sub-segment, dedicated risk and analytics platforms "
        "include MSCI RiskMetrics, BlackRock Aladdin, and Bloomberg PORT. These tools focus "
        "primarily on portfolio-level risk attribution and are not designed for cross-asset "
        "regime-based correlation monitoring. Open-source alternatives - QuantLib, zipline, bt "
        "- address backtesting and portfolio optimisation but lack integrated visualisation, "
        "real-time data ingestion, and AI synthesis layers.")

    add_space(doc, 4)
    add_table_caption(doc, "Table 1.1 - Major Financial Analytics Platforms: Feature Comparison")
    add_styled_table(doc, [
        ["Platform", "Cross-Asset Correlation", "Regime Detection", "AI Synthesis", "Geo Monitoring", "Annual Cost (USD)"],
        ["Bloomberg Terminal", "Yes", "Limited", "Bloomberg GPT (beta)", "Manual", "~$25,000"],
        ["Refinitiv Eikon", "Yes", "None", "News NLP only", "Manual", "~$22,000"],
        ["FactSet", "Partial", "None", "FactSet Signals", "None", "~$18,000"],
        ["MSCI RiskMetrics", "Yes", "Limited", "None", "None", "~$30,000+"],
        ["This Platform", "Yes (DCC-GARCH)", "Yes (Markov)", "7-Agent Pipeline", "Yes (Automated)", "Open-source"],
    ], [4.5, 2.8, 2.6, 3.0, 2.6, 2.6])
    add_para(doc, "Source: Publicly available vendor pricing and feature documentation, 2025.",
             italic=True, color=MUTED, size=8, space_after=10)

    add_sub_heading(doc, "1.2  Cross-Asset Spillover Methods: Academic State of the Art")
    add_para(doc,
        "The academic literature on cross-asset spillover analysis has produced several "
        "foundational methodologies that inform this project. The key contributions "
        "are summarised in the table below.")
    add_space(doc, 4)
    add_table_caption(doc, "Table 1.2 - Foundational Spillover Methodologies")
    add_styled_table(doc, [
        ["Method", "Authors", "Year", "Application in This Project"],
        ["Granger Causality", "C.W.J. Granger", "1969", "Directional lead-lag identification between asset pairs at 5% significance"],
        ["Markov Regime-Switching", "J.D. Hamilton", "1989", "Four-state correlation regime classification with forward transition forecasts"],
        ["GARCH / DCC-GARCH", "R.F. Engle", "1982 / 2002", "Dynamic conditional correlation capturing volatility clustering"],
        ["FEVD Connectedness", "Diebold & Yilmaz", "2009–2014", "System-wide spillover index and directional net transmitter scores"],
        ["Transfer Entropy", "T. Schreiber", "2000", "Nonlinear information flow beyond the linear Granger framework"],
        ["Early Warning Systems", "Kaminsky et al.", "1998", "Composite signal-based regime transition detection"],
    ], [3.8, 3.2, 1.8, 9.3])
    add_para(doc, "Source: Original publications. Python implementations via statsmodels, arch, scipy.",
             italic=True, color=MUTED, size=8, space_after=10)

    add_para(doc,
        "While each of these methods is well-established individually, their integration into "
        "a single, accessible, real-time platform with a unified interface remains absent from "
        "both commercial and open-source offerings. The academic-practitioner gap - where "
        "rigorous econometric methods remain confined to research code and are not "
        "operationalised for day-to-day market monitoring - represents the primary intellectual "
        "motivation for this project.")

    add_sub_heading(doc, "1.3  AI in Financial Analytics")
    add_para(doc,
        "The application of large language models (LLMs) to financial analysis is nascent but "
        "rapidly maturing. Bloomberg introduced BloombergGPT in 2023 - a 50-billion parameter "
        "model trained on financial text corpora, focused on information retrieval and "
        "summarisation. Kensho Technologies (acquired by S&P Global) pioneered NLP-based event "
        "analysis. Refinitiv News Analytics provides automated sentiment scoring across news "
        "feeds. FactSet has introduced Signals for quantitative factor generation.")
    add_para(doc,
        "However, current AI financial tools share a critical limitation: they operate as "
        "single-agent, single-domain tools performing text processing or data summarisation. "
        "None implement multi-agent orchestrated workflows where specialised agents share "
        "structured outputs, build on each other's analysis, and flag inter-agent "
        "disagreements. This multi-agent synthesis architecture is the primary AI innovation "
        "in this project.")

    add_callout(doc,
        "KEY OBSERVATION",
        "Existing AI financial tools summarise data. This platform synthesises it - through "
        "seven agents operating in three dependency-ordered rounds, each informed by the "
        "structured outputs of its upstream peers before generating its own analysis.",
        bg_hex=_CARD_BG, border_hex="C9A84C")

    add_sub_heading(doc, "1.4  Competitive Landscape")
    add_para(doc,
        "Beyond institutional platforms, the landscape includes data visualisation and fintech "
        "analytics tools targeting retail investors and smaller institutions: TradingView "
        "(technical analysis, basic correlation), Koyfin (cross-asset charting), Portfolio "
        "Visualizer (historical backtesting), Macroaxis and Seeking Alpha (AI-assisted "
        "fundamental analysis). None of these platforms implement quantitative spillover "
        "analytics, geopolitical event-window forensics, parametric scenario propagation "
        "with live betas, or multi-agent AI synthesis.")


def build_section2(doc):
    doc.add_page_break()
    add_section_heading(doc, "2", "Gap Analysis",
        "Identification of unaddressed market needs and structural limitations in existing platforms.")

    add_sub_heading(doc, "2.1  Identified Gaps in Existing Solutions")

    gaps = [
        ("Gap 1: Cost Accessibility",
         "Bloomberg Terminal and Refinitiv Eikon are priced at USD 18,000–30,000 per annum. "
         "This structure is inaccessible to academic researchers, emerging market practitioners, "
         "independent analysts, and smaller investment funds. No open-source alternative "
         "provides comparable cross-asset analytical breadth with a unified interface."),
        ("Gap 2: Integration Fragmentation",
         "Spillover analytics, geopolitical risk monitoring, macro intelligence, implied "
         "volatility tracking, and trade idea generation currently exist as separate, "
         "disconnected workflows. A practitioner must query multiple platforms and manually "
         "synthesise outputs - introducing latency and human error in high-volatility "
         "environments. No unified, end-to-end workflow exists."),
        ("Gap 3: Static Correlation Methods",
         "Standard rolling correlation (Pearson) uses fixed lookback windows and does not "
         "account for volatility clustering. DCC-GARCH - methodologically superior for "
         "regime-change environments - is absent from all publicly available non-terminal "
         "dashboards."),
        ("Gap 4: No Forward-Looking Regime Forecasting",
         "Existing tools classify the current regime but do not provide probabilistic "
         "forecasts of regime transitions. The Markov transition matrix, steady-state "
         "distribution, and mean first-passage time to Crisis - outputs directly informing "
         "risk management decisions - are absent from standard analytics platforms."),
        ("Gap 5: AI Synthesis vs. AI Summarisation",
         "Current AI financial tools retrieve and summarise data. They do not perform "
         "multi-domain synthesis across simultaneously active analytical frameworks. Seven "
         "specialist agents - macro, geopolitical, risk, commodities, stress, trade, signal - "
         "each reading its peers' structured outputs before forming its own view is "
         "qualitatively distinct from any existing commercial offering."),
        ("Gap 6: Geopolitical-Market Linkage",
         "No publicly available tool systematically maps geopolitical shocks to cross-asset "
         "correlation regime changes with empirical event-window analysis. Forensic examination "
         "of pre, during, and post-event performance across thirteen shocks since 2008 - with "
         "vol shift and correlation regime attribution - is not available on any "
         "accessible platform."),
        ("Gap 7: Parametric Scenario Propagation",
         "Scenario analysis tools are either proprietary or overly simplified. No accessible "
         "tool combines parametric multi-channel shock inputs (oil, gold, yields, DXY, credit, "
         "geopolitical) with live OLS beta propagation to a full cross-asset universe and "
         "outputs VaR and ES at multiple confidence levels."),
    ]

    for title, body_text in gaps:
        add_mixed_para(doc,
            [(title + " - ", True, False, GOLD, 9.5),
             (body_text, False, False, BODY, 9.5)],
            space_after=7)

    add_sub_heading(doc, "2.2  Opportunity Summary")
    add_para(doc,
        "The convergence of three trends creates a timely opportunity: (1) the increasing "
        "availability of institutional-grade free data via Yahoo Finance and FRED; (2) the "
        "maturation of open-source econometric libraries implementing rigorous academic methods; "
        "and (3) the emergence of accessible, high-capability LLM APIs enabling multi-agent "
        "orchestration. The opportunity is to build an institutional-grade cross-asset analytics "
        "platform that is fully open-source, deployable at zero marginal cost, and "
        "methodologically rigorous.")

    add_space(doc, 4)
    add_table_caption(doc, "Table 2.1 - Opportunity Matrix")
    add_styled_table(doc, [
        ["Gap Identified", "Opportunity Addressed", "Priority"],
        ["Cost accessibility", "Open-source, free-to-deploy platform", "Critical"],
        ["Integration fragmentation", "Unified 14-page analytical workspace", "Critical"],
        ["Static correlation only", "DCC-GARCH dynamic conditional correlation", "High"],
        ["No regime forecasting", "Markov transition matrix with first-passage times", "High"],
        ["AI summarisation only", "7-agent dependency-ordered synthesis pipeline", "High"],
        ["No geo-market forensics", "13-event window analysis with regime attribution", "Medium"],
        ["No scenario propagation", "6-channel parametric shock engine with live betas", "High"],
    ], [5.5, 8.0, 2.6])


def build_section3(doc):
    doc.add_page_break()
    add_section_heading(doc, "3", "Project Definition and Scope",
        "Formal statement of the problem, system objectives, coverage universe, "
        "analytical architecture, and scope boundaries.")

    add_sub_heading(doc, "3.1  Problem Statement")
    add_callout(doc,
        "PROBLEM STATEMENT",
        "Portfolio managers, risk analysts, and researchers require a unified, real-time view "
        "of cross-asset stress transmission dynamics - integrating correlation regime detection, "
        "directional spillover identification, geopolitical event monitoring, macro intelligence, "
        "and AI-synthesised trade ideas - at a cost accessible beyond institutional terminal "
        "licensing. No such platform currently exists in the open-source or low-cost "
        "commercial landscape.",
        bg_hex=_CARD_BG, border_hex="C9A84C")

    add_sub_heading(doc, "3.2  Objectives")
    add_para(doc, "The Equity-Commodities Spillover Monitor is designed to achieve the "
             "following primary objectives:", space_after=4)
    objectives = [
        "Provide real-time correlation regime classification (four states: Decorrelated, Normal, "
        "Elevated, Crisis) with Markov-based forward transition probability forecasts.",
        "Implement institutional-grade spillover methods - Granger causality, transfer entropy, "
        "Diebold-Yilmaz FEVD - in an accessible, interactive interface.",
        "Enable parametric scenario analysis via multi-channel shock propagation using live "
        "OLS betas, with VaR and ES tail risk quantification.",
        "Provide forensic event-window analysis across thirteen tracked geopolitical shocks "
        "since the Global Financial Crisis.",
        "Deploy a seven-agent AI orchestration pipeline synthesising macro, geopolitical, "
        "commodity, risk, stress, and trade intelligence in a structured dependency graph.",
        "Generate a composite Early Warning System (EWS) score identifying conditions that "
        "historically precede correlation regime transitions.",
        "Serve as an open-source, educationally accessible alternative to institutional-grade "
        "cross-asset analytics platforms.",
    ]
    for obj in objectives:
        add_bullet(doc, obj)
    add_space(doc, 6)

    add_sub_heading(doc, "3.3  Coverage Universe")
    add_space(doc, 4)
    add_table_caption(doc, "Table 3.1 - Coverage Universe")
    add_styled_table(doc, [
        ["Asset Class", "Count", "Instruments"],
        ["Global Equity Indices", "15",
         "S&P 500, Nasdaq 100, DJIA, Russell 2000 (USA); Eurostoxx 50, DAX, CAC 40, "
         "FTSE 100 (Europe); Nikkei 225, TOPIX (Japan); Hang Seng, Shanghai Composite, "
         "CSI 300 (China); Sensex, Nifty 50 (India)"],
        ["Commodity Futures", "17",
         "WTI Crude, Brent Crude, Natural Gas, RBOB Gasoline, Heating Oil (Energy); "
         "Gold, Silver, Platinum (Precious Metals); Copper, Aluminum, Nickel (Industrial); "
         "Wheat, Corn, Soybeans, Sugar #11, Coffee, Cotton (Agriculture)"],
        ["Fixed Income ETFs", "6",
         "TLT (20Y+ Treasury), LQD (IG Corporate), HYG (HY Corporate), "
         "EMB (EM USD Bonds), SHY (Short Treasury), TIP (TIPS)"],
        ["FX Pairs", "6",
         "DXY Dollar Index, EUR/USD, GBP/USD, USD/JPY, USD/CNY, USD/INR"],
        ["Implied Volatility", "4",
         "VIX (CBOE Equity), OVX (CBOE Oil), GVZ (CBOE Gold), VVIX (Volatility of VIX)"],
        ["Geopolitical Events", "13",
         "GFC (2008), Arab Spring (2011), US-China Trade War (2018), Aramco Attack (2019), "
         "COVID-19 (2020), WTI Negative (2020), Ukraine War (2022), LME Nickel (2022), "
         "Fed Hike Cycle (2022), SVB Crisis (2023), Israel-Hamas (2023), "
         "India-Pakistan (2025), Iran-Hormuz (2025)"],
    ], [4.0, 1.4, 12.7])

    add_sub_heading(doc, "3.4  Scope Boundaries")
    add_para(doc, "In scope:", bold=True, size=9.5, space_after=3)
    in_scope = [
        "Daily close price data for all equity, commodity, fixed income, and FX instruments",
        "Live macro data via FRED API (yield curve, CPI, Fed Funds Rate, GDP, PMI)",
        "RSS-based automated geopolitical headline ingestion and severity scoring",
        "Correlation regime detection, DCC-GARCH, and Markov regime forecasting",
        "Granger causality, transfer entropy, and Diebold-Yilmaz FEVD spillover analytics",
        "Parametric scenario propagation using live OLS betas",
        "Historical portfolio stress testing against all thirteen tracked events",
        "Seven-agent AI orchestration pipeline with dependency-ordered peer context sharing",
        "Maritime chokepoint disruption monitoring for six critical straits",
    ]
    for item in in_scope:
        add_bullet(doc, item)
    add_space(doc, 4)
    add_para(doc, "Out of scope:", bold=True, size=9.5, space_after=3)
    out_scope = [
        "Order execution or trading infrastructure of any kind",
        "Real-time tick data at sub-daily resolution",
        "Options pricing, Greeks, or volatility surface analytics",
        "Fundamental analysis, earnings models, or company-level financial data",
        "Credit default swap pricing or single-name credit analytics",
        "Regulatory reporting or compliance workflows",
    ]
    for item in out_scope:
        add_bullet(doc, item)


def build_section4(doc):
    doc.add_page_break()
    add_section_heading(doc, "4", "Technology Choices and Justification",
        "Full-stack technology selection rationale across data, analytics, AI, and deployment layers.")

    add_sub_heading(doc, "4.1  Technology Stack Overview")
    add_space(doc, 4)
    add_table_caption(doc, "Table 4.1 - Full Technology Stack")
    add_styled_table(doc, [
        ["Layer", "Library / Service", "Version", "Primary Role"],
        ["UI Framework", "Streamlit", "1.x", "Multi-page web application, interactive widgets, data caching"],
        ["Visualisation", "Plotly", "5.x", "Interactive charts, heatmaps, network graphs, choropleths"],
        ["Market Data", "yfinance", "0.2.x", "Historical and live equity, commodity, ETF, and FX data"],
        ["Macro Data", "fredapi", "0.5.x", "FRED macro series: yield curve, CPI, GDP, PMI, Fed Funds Rate"],
        ["News / RSS", "feedparser", "6.x", "Automated geopolitical RSS ingestion from 7 major sources"],
        ["Time Series", "pandas / numpy", "2.x / 1.x", "Data wrangling, rolling statistics, return computation"],
        ["Econometrics", "statsmodels", "0.14.x", "Granger causality, VAR, OLS regression, regime statistics"],
        ["GARCH Models", "arch", "6.x", "DCC-GARCH dynamic conditional correlation estimation"],
        ["ML / Validation", "scikit-learn", "1.x", "Regime classification, walk-forward cross-validation"],
        ["Scientific", "scipy", "1.x", "Transfer entropy, statistical testing, signal processing"],
        ["Graph Analytics", "networkx", "3.x", "Spillover network construction and layout algorithms"],
        ["AI - Primary", "anthropic", "0.x", "Claude Sonnet: 7-agent orchestration pipeline and AI Analyst"],
        ["AI - Fallback", "openai", "1.x", "GPT-4o: fallback provider for all AI agent functions"],
        ["PDF Generation", "reportlab", "4.x", "Automated PDF report export from Python"],
        ["Image Processing", "Pillow", "10.x", "Image asset handling for generated reports"],
        ["Deployment", "Render", "-", "Cloud hosting with Secret File support for Streamlit secrets"],
    ], [2.8, 3.0, 2.0, 10.3])

    add_sub_heading(doc, "4.2  Justification by Layer")

    justifications = [
        ("Streamlit (UI Framework)",
         "Streamlit was selected over Flask, FastAPI, or Dash for three reasons: (1) it is "
         "Python-native, eliminating the need for JavaScript development; (2) its caching "
         "decorators (@st.cache_data, @st.cache_resource) provide built-in TTL-based data "
         "freshness management critical for a live dashboard; and (3) it supports free-tier "
         "cloud deployment with minimal configuration. The multi-page architecture and session "
         "state management are sufficient for the dashboard's complexity without introducing "
         "unnecessary framework overhead."),
        ("yfinance + fredapi (Data Layer)",
         "yfinance provides free, reliable access to Yahoo Finance historical and live data "
         "across equities, commodity ETFs, FX pairs, and volatility indices with no API key "
         "requirement for basic usage. The fredapi provides official Federal Reserve Economic "
         "Data access for all macro series. Together, these two libraries cover the full data "
         "universe without licensing costs. The primary limitation - absence of true commodity "
         "futures tick data - is acceptable given the daily-resolution analytical focus."),
        ("statsmodels + arch (Econometric Layer)",
         "statsmodels is the standard Python implementation of econometric methods with "
         "peer-reviewed documentation and academic acceptance. It provides Granger causality "
         "tests, VAR, and OLS regression. The arch library is the canonical Python "
         "implementation of GARCH-family models, including DCC-GARCH as specified by "
         "Engle (2002). Both libraries are actively maintained and used in published "
         "academic research."),
        ("anthropic / Claude Sonnet (AI Primary)",
         "Claude Sonnet was selected as the primary AI provider for three reasons: "
         "(1) the 200K-token context window supports full dashboard state injection into "
         "every agent query; (2) Claude's instruction-following and structured output quality "
         "is consistently superior for financial reasoning tasks; and (3) Anthropic's API "
         "pricing is competitive relative to GPT-4o for the volume of agent calls per "
         "session. OpenAI GPT-4o serves as a fallback provider."),
        ("Render (Deployment)",
         "Render was selected over Streamlit Community Cloud and Heroku because it supports "
         "Secret Files - a mechanism for placing actual files (including .streamlit/secrets.toml) "
         "on the server. Streamlit's st.secrets mechanism reads from the file system, not "
         "environment variables. Render's free tier is sufficient for demonstration "
         "and academic use."),
    ]

    for title, body_text in justifications:
        add_mixed_para(doc,
            [(title + " - ", True, False, GOLD, 9.5),
             (body_text, False, False, BODY, 9.5)],
            space_after=8)


# ════════════════════════════════════════════════════════════════════════════
# PART II
# ════════════════════════════════════════════════════════════════════════════
def build_part2_intro(doc):
    add_part_header(doc, "Part II", "User Manual")
    add_para(doc,
        "This user manual provides comprehensive operating guidance for the Equity-Commodities "
        "Spillover Monitor. It covers system setup, navigation, a page-by-page operating guide, "
        "AI workforce reference, output interpretation, and troubleshooting.",
        size=10, space_after=8)


def build_section5(doc):
    add_section_heading(doc, "5", "Getting Started",
        "System requirements, installation, configuration, and navigation overview.")

    add_sub_heading(doc, "5.1  System Requirements and Installation")
    add_para(doc, "Requirements: Python 3.10 or later. All dependencies are listed in requirements.txt.",
             space_after=4)
    steps = [
        "Clone the repository: git clone https://github.com/HPATKAR/equity-commodities-spillover.git",
        "Navigate to the project directory: cd equity-commodities-spillover",
        "Install all dependencies: pip install -r requirements.txt",
        "Launch the application: streamlit run app.py",
        "Open a browser and navigate to http://localhost:8501",
    ]
    for step in steps:
        add_bullet(doc, step)
    add_space(doc, 4)
    add_para(doc,
        "The application will load with live data on first run. Data is cached with TTL values "
        "ranging from 15 minutes (implied volatility) to 1 hour (correlation matrices, agent "
        "outputs). Initial load on a fresh environment may take 30–60 seconds as all data "
        "is fetched and models are computed.")

    add_sub_heading(doc, "5.2  Configuration (API Keys)")
    add_para(doc,
        "API keys are managed via .streamlit/secrets.toml (local) or Render Secret Files "
        "(deployed). The application functions without API keys but with reduced capability: "
        "AI agent outputs and FRED Macro Dashboard data will not be available.")
    add_space(doc, 4)
    add_table_caption(doc, "Table 5.1 - API Key Configuration")
    add_styled_table(doc, [
        ["Key", "Provider", "Required", "Purpose"],
        ["anthropic_api_key", "console.anthropic.com", "Recommended",
         "All 7 AI agents + AI Analyst chatbot (preferred provider)"],
        ["openai_api_key", "platform.openai.com", "Optional",
         "Fallback provider for all AI agents and chatbot"],
        ["fred_api_key", "fred.stlouisfed.org/docs/api", "Optional",
         "Macro Dashboard yield curve, CPI, GDP, and PMI data"],
    ], [3.8, 4.0, 2.4, 8.0])
    add_callout(doc,
        "SECRETS.TOML FORMAT",
        "[keys]\n"
        "anthropic_api_key = \"sk-ant-...\"\n"
        "openai_api_key    = \"sk-proj-...\"\n"
        "fred_api_key      = \"your_fred_key\"\n\n"
        "[config]\n"
        "enable_auth      = false\n"
        "refresh_interval = 300",
        bg_hex=_CARD_BG, border_hex="3A9B68")
    add_para(doc,
        "Important - Render deployment: API keys must be added as a Secret File at path "
        ".streamlit/secrets.toml, not as Environment Variables. Streamlit's st.secrets "
        "mechanism reads from the file system, not environment variables.",
        italic=True, color=MUTED, size=9)

    add_sub_heading(doc, "5.3  Navigation Overview")
    add_space(doc, 4)
    add_table_caption(doc, "Table 5.2 - Navigation Structure")
    add_styled_table(doc, [
        ["Navigation Group", "Pages Included"],
        ["(Landing)", "Overview"],
        ["Macro", "Macro Intelligence Dashboard"],
        ["Analysis", "War Impact Map  ·  Geopolitical Triggers  ·  Correlation Analysis  ·  Spillover Analytics"],
        ["Strategy", "Trade Ideas  ·  Portfolio Stress Test  ·  Scenario Engine"],
        ["Monitor", "Commodities to Watch  ·  Strait Watch"],
        ["Research", "Performance Review  ·  AI Analyst"],
        ["Insights", "Insights"],
        ["About", "Heramb S. Patkar  ·  Jiahe Miao  ·  Ilian Zalomai"],
    ], [4.5, 13.6])


def build_section6(doc):
    doc.add_page_break()
    add_section_heading(doc, "6", "Page-by-Page Operating Guide",
        "Detailed description of each page: purpose, key components, controls, and interpretation guidance.")

    pages = [
        {
            "name": "Overview",
            "purpose": "The primary dashboard landing page. Provides a complete cross-asset market state summary in a single view.",
            "components": [
                ("KPI Strip", "Five live metrics: Correlation Regime (color-coded), 60-day average cross-asset correlation with 1-month delta, Best Equity (1M return), Worst Equity (1M), Best Commodity (1M)."),
                ("FI/FX Context Strip", "Five fixed income and FX signals: TLT 30-day return (duration), HYG vs. LQD spread (credit stress), S&P 500 vs. TLT 60-day correlation (hedge effectiveness), DXY 30-day change, duration slope (TLT vs. SHY)."),
                ("Implied Volatility Panel", "Expandable section showing live OVX (oil vol), GVZ (gold vol), VIX (equity vol), VVIX (vol-of-vol) with freshness indicators and historical charts."),
                ("Risk Gauge", "Composite 0–100 score across five weighted components. Green zone: 0–30. Amber zone: 30–55. Red zone: 55–100. Updates every 5 minutes."),
                ("Early Warning System", "Five signal cards (correlation slope, vol regime, cross-asset stress, macro momentum, geopolitical escalation) with composite EWS score and historical analogue table."),
                ("Equity-Commodity Heatmap", "Pairwise correlation heatmap across all equities and commodities. Color intensity indicates correlation strength."),
                ("AI Workforce Output", "Risk Officer morning briefing displayed inline. All seven agents accessible in expandable section. Divergence flags highlighted when agents reach conflicting conclusions."),
            ],
            "interpretation": "Begin every session on this page. The regime label and risk score provide immediate context for all downstream analysis. An EWS score above 70 warrants heightened monitoring across the analysis pages.",
        },
        {
            "name": "Macro Intelligence Dashboard",
            "purpose": "FRED-powered macro context integrated with the AI Macro Strategist agent.",
            "components": [
                ("Yield Curve Chart", "10Y minus 2Y Treasury spread with recession shading and inversion duration tracker."),
                ("CPI YoY", "Consumer Price Index year-over-year change - key input to the Macro Strategist's inflation assessment."),
                ("Fed Funds Rate", "Effective federal funds rate with rate cycle annotations."),
                ("GDP Growth", "Real GDP growth quarter-over-quarter with recession shading."),
                ("ISM PMI", "ISM Manufacturing PMI. Readings below 50 indicate contraction."),
                ("AI Macro Strategist", "Agent briefing incorporating yield curve shape, inflation posture, Fed stance, and peer context from the Geopolitical Analyst and Risk Officer."),
            ],
            "interpretation": "Cross-reference the yield curve shape with the current correlation regime. Historically, inverted yield curves coincide with elevated equity-bond correlation (hedge breakdown). Read alongside the Risk Officer's brief on Overview.",
        },
        {
            "name": "War Impact Map",
            "purpose": "Global choropleth quantifying equity market exposure across three active conflict theaters.",
            "components": [
                ("Choropleth Map", "195 countries scored by equity exposure. Toggle between Combined, Ukraine, Israel-Hamas, and Iran/Hormuz views."),
                ("Top 25 Exposed Markets", "Ranked table with per-conflict score columns and actual equity index performance since each conflict's onset date."),
            ],
            "interpretation": "High exposure score plus negative realized performance indicates the market has priced the risk. High exposure score plus stable performance indicates potential unpriced risk - a positioning opportunity.",
        },
        {
            "name": "Geopolitical Triggers",
            "purpose": "Forensic event-window analysis for all thirteen tracked geopolitical shocks since 2008.",
            "components": [
                ("Event Selector", "Dropdown selection from thirteen tracked events with date range and description."),
                ("Lookback Controls", "Pre-event window (15–90 days) and post-event window (15–180 days) sliders. Asset multiselect for customising the comparison set."),
                ("Indexed Performance Chart", "All selected assets indexed to 100 at event onset. Vertical band marks the event window."),
                ("Volatility Shift Chart", "Pre- vs. post-event annualized realized volatility grouped by asset. Persistent post-event vol elevation signals ongoing disruption."),
                ("Correlation Shift Table", "Asset pairs sorted by correlation shift magnitude. Red shift indicates contagion; green shift indicates decoupling."),
                ("Live Intelligence Feed", "Automated RSS ingestion showing current geopolitical headlines tagged by region, commodity, and severity score."),
            ],
            "interpretation": "When analysing a current event, select the most analogous historical event and compare the pre-event window charts. Correlation regime changes at event onset are the most reliable early indicator of whether a shock is systemic or contained.",
        },
        {
            "name": "Correlation Analysis",
            "purpose": "Four-tab quantitative correlation analysis: rolling Pearson, DCC-GARCH, regime detection, and Markov regime forecast.",
            "components": [
                ("Rolling Correlation Tab", "Select any two assets and a lookback window (21–252 days). Multi-pair overlay expander shows four default pairs simultaneously."),
                ("DCC-GARCH Tab", "Dynamic Conditional Correlation overlaid against rolling Pearson. Divergence indicates vol-regime-dependent correlation behaviour."),
                ("Regime Detection Tab", "Scatter of all trading days colored by four-state regime. Compact table shows time spent in each regime."),
                ("Markov Forecast Tab", "Transition probability heatmap, steady-state distribution, mean days to Crisis from each state, and average run length per regime."),
            ],
            "interpretation": "DCC-GARCH value diverging above rolling Pearson indicates vol-regime-driven correlation elevation - a leading indicator of potential regime transition. Mean days to Crisis is a probabilistic estimate, not a precise prediction.",
        },
        {
            "name": "Spillover Analytics",
            "purpose": "Three-method cross-asset spillover identification: Granger causality, transfer entropy, and Diebold-Yilmaz FEVD.",
            "components": [
                ("Granger Causality Heatmap", "P-value matrix for commodity-to-equity causality. Red cells (p < 0.05) indicate statistically significant predictive relationships. Top pairs table sorted by p-value."),
                ("Transfer Entropy Matrix", "Net information flow from commodities to equities capturing nonlinear dependencies beyond linear Granger tests."),
                ("Diebold-Yilmaz FEVD Heatmap", "Forecast Error Variance Decomposition. Total Spillover Index summarises system-wide connectedness on a 0–100% scale."),
                ("Network Graph", "Node size equals net transmitter score. Edge width equals spillover strength. Green nodes transmit; red nodes absorb."),
                ("Cross-Asset Section", "Granger causality table for rates, FX, and fixed income channels (S&P 500, Gold, WTI, TLT, HYG, DXY)."),
            ],
            "interpretation": "Assets with large positive net transmitter scores are the primary shock sources. When a Granger-significant transmitter commodity makes a significant move, the linked equity index has historically followed within the identified lag period.",
        },
        {
            "name": "Trade Ideas",
            "purpose": "Regime-triggered cross-asset trade cards with AI Trade Structurer output.",
            "components": [
                ("Regime Filter", "Cards automatically filtered to ideas valid for the current regime. 'Show all regimes' checkbox overrides the filter."),
                ("Trade Cards", "Each card shows: directional thesis (Long/Short), category and regime trigger, rationale paragraph, entry signal, exit conditions, risk factors, and a live 60-day rolling correlation chart."),
                ("AI Trade Structurer", "Generates new trade ideas incorporating full peer context from Macro Strategist, Geopolitical Analyst, Stress Engineer, and Commodities Specialist."),
                ("PDF Report", "One-click PDF generation of all active trade cards with regime context."),
            ],
            "interpretation": "Use the live correlation chart on each card to confirm the pair relationship underpinning the trade is currently intact. A correlation that has moved significantly from the expected range is a signal to reassess the trade rationale before entry.",
        },
        {
            "name": "Portfolio Stress Test",
            "purpose": "Custom portfolio construction and stress testing against all thirteen tracked historical events.",
            "components": [
                ("Asset Selection", "Preset allocations or custom builds from global indices, commodities, fixed income ETFs, and individual S&P 500 stocks. Free-form ticker entry supported."),
                ("Weight Editor", "Per-asset weight inputs auto-normalised to 100%."),
                ("Stress Test Results", "Per-event table: Pre-Event Return, During Return, Post-Event Return, Max Drawdown, Sharpe Ratio. Color-coded by return sign."),
                ("Summary Statistics", "Average return across all events, historical VaR at the 5th percentile, and win rate."),
            ],
            "interpretation": "The worst 5% return is the most operationally relevant metric for risk committee conversations. Compare the portfolio's worst event return against its average return to assess the asymmetry of downside risk.",
        },
        {
            "name": "Scenario Engine",
            "purpose": "Forward-looking parametric shock propagation with VaR and ES outputs.",
            "components": [
                ("Preset Scenarios", "Six presets: Oil Supply Shock (+40%), Risk-Off Flight, Rate Shock (+150bps), Strait Closure, China Hard Landing, Stagflation."),
                ("Shock Sliders", "Six inputs: Oil (%), Gold (%), Natural Gas (%), Copper (%), Yield Shift (bps), DXY (%), Credit Spreads (bps), Geopolitical Factor (0–10)."),
                ("Impact Charts", "Estimated return impact on 10 equity indices and 8 commodity futures via OLS beta propagation."),
                ("VaR / ES Table", "Historical and shocked VaR at 95% and 99% confidence. Expected Shortfall at 95% and 99%."),
            ],
            "interpretation": "The OLS beta propagation model uses historical covariance. Results should be treated as order-of-magnitude estimates. ES is the more conservative measure - it captures average loss in the tail beyond VaR.",
        },
        {
            "name": "Commodities to Watch",
            "purpose": "Live commodity watchlist with intraday, daily, and COT positioning views.",
            "components": [
                ("Live Snapshot Table", "All 17 futures: price, 1D%, 5D%, YTD%, annualised vol, regime label."),
                ("Intraday Charts", "Selected commodities indexed to 100 at lookback start. Rolling 24-hour annualised vol chart. Elevated threshold: 40% annualised."),
                ("Daily Chart", "Indexed daily prices with geopolitical event bands overlaid. Rolling 30-day vol."),
                ("COT Positioning", "Net speculative positioning vs. commercial hedgers from CFTC Commitments of Traders data. Extremes table flags crowded positioning above +25% or below -25% of open interest."),
            ],
            "interpretation": "The COT extreme signal is a contrarian indicator - extreme speculative longs historically precede corrections; extreme shorts precede rallies. Most reliable in Elevated and Crisis regimes.",
        },
        {
            "name": "Strait Watch",
            "purpose": "Real-time disruption monitoring for six critical maritime chokepoints.",
            "components": [
                ("Global KPI Strip", "Total oil at risk (mb/d), global percentage at risk, worst strait name, Critical/Elevated count."),
                ("Strait Cards", "Per-strait: disruption score (0–100), vessel traffic vs. baseline, risk factors, commodity channels at risk, estimated daily trade value at risk."),
                ("Status Levels", "NORMAL (0–25), CAUTION (25–50), ELEVATED (50–75), CRITICAL (75–100)."),
                ("Commodity Price Context", "Brent crude price with disruption events marked. WTI-Brent spread as supply route risk premium indicator."),
            ],
            "interpretation": "CRITICAL status on Hormuz has historically correlated with Brent price spikes within 5–10 trading days. The WTI-Brent spread widening is a real-time market signal of perceived supply route risk. Cross-reference with OVX on Overview.",
        },
        {
            "name": "Performance Review",
            "purpose": "Model validation and signal accuracy metrics, all computed out-of-sample.",
            "components": [
                ("Regime Detection", "Confusion matrix, balanced accuracy, precision, recall, F1 score, and AUC for rule-based and ML walk-forward classifiers."),
                ("Granger Hit Rates", "Directional accuracy per significant asset pair. 'Edge' column shows hit rate minus 50% (positive edge = predictive value)."),
                ("EWS Lead-Lag", "Correlation between composite risk score and VIX at lags of 0–20 days. Peak lag identifies the average lead time of the EWS over realised equity volatility."),
                ("COT Contrarian Signals", "Win rate, mean reversion days, and average drawdown per commodity for COT extreme positioning signals."),
            ],
            "interpretation": "Only Granger pairs with positive edge are used to generate trade ideas. The EWS lead-lag peak (typically 5–7 days) provides the operating horizon for acting on elevated early warning scores.",
        },
        {
            "name": "AI Analyst",
            "purpose": "Claude Sonnet / GPT-4o chatbot with full live dashboard state injected into every query.",
            "components": [
                ("Context Status Bar", "Shows age of injected market context. 'Refresh context' button forces a live update."),
                ("Suggested Questions", "Ten pre-built queries covering regime interpretation, spillover identification, hedge effectiveness, and scenario analysis."),
                ("Chat Interface", "Free-form text input with real-time response streaming. Full conversation history maintained within session."),
                ("Context Inspector", "Expandable panel showing the exact market state text injected into every query for full transparency."),
                ("Agent Activity Log", "Structured audit trail of every agent action, routing decision, and escalation from the current session."),
            ],
            "interpretation": "The AI Analyst's responses are grounded in injected dashboard state, not training data. The context inspector allows verification of what data the model used for any given response.",
        },
        {
            "name": "Insights",
            "purpose": "Plain-language verdicts synthesising all analytical layers into actionable one-sentence directives.",
            "components": [
                ("Overall Stress Card", "Risk score interpretation with action directive and confidence score."),
                ("Diversification Status Card", "Regime-based hedge effectiveness. Green = working. Red = broken."),
                ("Leading Commodity Card", "Commodity with the highest current absolute correlation to the equity average."),
                ("Early Warning Card", "EWS composite score with regime transition probability context."),
                ("Private Credit Risk Card", "HY OAS, BKLN, BDC, CDX HY composite signal."),
                ("India / Rupee Card", "Oil dependency, rupee transmission, and Nifty 50 driver assessment."),
            ],
            "interpretation": "Confidence below 60% indicates model uncertainty; treat as directional signal only. Use Insights as a starting point for deeper analysis on the relevant analytical page.",
        },
    ]

    for page in pages:
        add_subsub_heading(doc, page["name"])
        add_para(doc, page["purpose"], italic=True, color=MUTED, size=9.5, space_after=4)

        comp_data = [["Component", "Description"]]
        for comp_name, comp_desc in page["components"]:
            comp_data.append([comp_name, comp_desc])
        add_styled_table(doc, comp_data, [4.2, 13.9])

        add_callout(doc, "INTERPRETATION", page["interpretation"],
                    bg_hex=_CARD_BG, border_hex="C9A84C")
        add_space(doc, 4)


def build_section7(doc):
    doc.add_page_break()
    add_section_heading(doc, "7", "AI Workforce Reference",
        "Architecture, dependency structure, and output interpretation for all seven agents.")

    add_para(doc,
        "The AI Workforce consists of seven specialised agents organised into three sequential "
        "rounds. Agents within a round run in parallel; each round waits for the prior round "
        "to complete before starting. Every agent receives the structured outputs of its "
        "upstream peers before generating its own analysis.")
    add_space(doc, 4)
    add_table_caption(doc, "Table 7.1 - Agent Pipeline Architecture")
    add_styled_table(doc, [
        ["Round", "Agent", "Role", "Inputs Received"],
        ["1", "Signal Auditor", "Calibrates confidence scores from Granger hit rates", "Performance Review backtest data"],
        ["1", "Macro Strategist", "Yield curve, inflation, Fed posture, GDP context", "FRED macro data + Signal Auditor confidence"],
        ["1", "Geopolitical Analyst", "Active conflict risk, sanctions, strait disruption", "Active events + Strait Watch scores + RSS feed"],
        ["2", "Risk Officer", "Synthesises Round 1 into morning briefing", "Signal Auditor + Macro Strategist + Geopolitical Analyst"],
        ["2", "Commodities Specialist", "COT positioning, supply shocks, sector rotation", "Commodity data + Signal Auditor + Macro Strategist"],
        ["3", "Stress Engineer", "Scenario stress, tail risk, portfolio shock analysis", "All Round 1 outputs + Risk Officer + Commodities Specialist"],
        ["3", "Trade Structurer", "Regime-triggered trade ideas with full peer context", "All six previous agents' structured outputs"],
    ], [1.5, 3.6, 5.4, 7.6])
    add_space(doc, 4)

    add_para(doc,
        "Divergence Detection: The orchestrator automatically compares agent conclusions "
        "across four dimensions - macro direction, geopolitical risk level, commodity "
        "positioning, and overall risk posture. When agents reach materially conflicting "
        "conclusions, a divergence flag is generated and displayed on the Overview page.",
        space_after=6)
    add_para(doc,
        "Cache and Invalidation: Agent outputs are cached for one hour. If a correlation "
        "regime change is detected between cache writes, all agents are automatically "
        "invalidated and re-run on the next page load.")
    add_para(doc,
        "Chief Quality Officer (CQO): A separate CQO agent runs on each analytical page "
        "independently. It audits the page's methodology and outputs a structured note "
        "flagging assumption violations, data limitations, and model caveats.")


def build_section8(doc):
    doc.add_page_break()
    add_section_heading(doc, "8", "Interpreting Key Outputs",
        "Reference guide for the platform's most important quantitative outputs.")

    outputs = [
        ("Correlation Regime",
         "Four-state classification based on rolling 60-day average absolute pairwise "
         "correlation across all equity-commodity pairs.\n"
         "  •  Decorrelated (< 0.15): Low average correlation. Diversification benefits strong.\n"
         "  •  Normal (0.15–0.45): Standard market environment. Traditional allocation strategies functional.\n"
         "  •  Elevated (0.45–0.60): Above-average co-movement. Diversification weakening.\n"
         "  •  Crisis (> 0.60): High systemic correlation. Diversification largely eliminated."),
        ("Composite Risk Score (0–100)",
         "Weighted average of five components: Correlation Level, Volatility Regime, "
         "Early Warning Score, Macro Posture, and Geopolitical Stress. Scores below 30 "
         "represent low-stress environments; scores above 55 represent elevated-stress "
         "environments warranting risk reduction."),
        ("Diebold-Yilmaz Total Spillover Index",
         "Ranges from 0% to 100%. Represents the fraction of total forecast error variance "
         "attributable to cross-asset shocks. Values above 50% indicate a highly interconnected "
         "system. Values above 70% indicate systemic stress comparable to the 2008–2009 GFC."),
        ("Mean Days to Crisis (Markov)",
         "The expected number of trading days to first enter the Crisis regime from the current "
         "regime, computed as the mean first-passage time in the estimated Markov chain. "
         "This is a probabilistic expectation and should be treated as a planning horizon, "
         "not a precise prediction."),
        ("Value at Risk (VaR) and Expected Shortfall (ES)",
         "VaR at confidence level c is the loss threshold not exceeded with probability c. "
         "VaR 95% is the loss level exceeded on 5% of historical observation days. "
         "ES (also known as Conditional VaR or CVaR) is the expected loss given that the loss "
         "exceeds the VaR threshold - a coherent risk measure that captures the average "
         "magnitude of tail losses, not just their threshold."),
    ]

    for title, body_text in outputs:
        add_subsub_heading(doc, title)
        add_para(doc, body_text, space_after=8)


def build_section9(doc):
    doc.add_page_break()
    add_section_heading(doc, "9", "Troubleshooting and FAQ",
        "Common issues and their resolution.")

    add_styled_table(doc, [
        ["Issue", "Resolution"],
        ["AI agents show 'Add API key' message",
         "No API key is configured. Add anthropic_api_key (preferred) or openai_api_key to "
         ".streamlit/secrets.toml under the [keys] section. On Render, use a Secret File - "
         "not Environment Variables."],
        ["Macro Dashboard shows no data",
         "The FRED API key is missing or invalid. Add fred_api_key to .streamlit/secrets.toml. "
         "A free key is available at fred.stlouisfed.org."],
        ["yfinance returns empty data for a ticker",
         "Yahoo Finance may be temporarily rate-limiting. Wait 60 seconds and reload. "
         "If persistent, verify the ticker symbol at finance.yahoo.com."],
        ["DCC-GARCH model fails to converge",
         "Can occur with very short data windows or extreme return distributions. "
         "The platform falls back to rolling Pearson automatically. Increase the data "
         "lookback window or reduce the number of selected assets."],
        ["AI Workforce error on Overview page",
         "Check the error details expander for the full traceback. Common causes: "
         "API key invalid or rate-limited; data fetch failure; session state inconsistency. "
         "Refresh the page to reinitialise the agent pipeline."],
        ["Granger heatmap shows no significant pairs",
         "Increase the maximum lag parameter (default: 5 days) in the Spillover Analytics "
         "panel. Expand the asset selection - small sets reduce statistical power."],
        ["PDF report download not working",
         "Ensure reportlab and Pillow are installed: pip install reportlab Pillow."],
        ["Streamlit shows 'ScriptRunContext' error",
         "Thread-safety caching issue. Refresh the page. If persistent, clear the Streamlit "
         "cache from the top-right hamburger menu."],
    ], [5.5, 12.6])

    add_space(doc, 10)
    add_para(doc,
        "For questions or issue reports, please contact the development team via the GitHub "
        "repository at github.com/HPATKAR/equity-commodities-spillover.",
        italic=True, color=MUTED, size=9)

    # final rule
    add_space(doc, 16)
    p_fin = doc.add_paragraph()
    _add_bottom_border(p_fin, color="C9A84C", size=6)
    _set_para_spacing(p_fin, before=0, after=60)
    add_para(doc,
        "Equity-Commodities Spillover Monitor  ·  Purdue University  ·  "
        "Mitchell E. Daniels, Jr. School of Business  ·  MGMT 69000-120: AI for Finance  ·  "
        "Heramb S. Patkar  ·  Jiahe Miao  ·  Ilian Zalomai  ·  April 2026",
        color=MUTED, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)


# ════════════════════════════════════════════════════════════════════════════
# MAIN
# ════════════════════════════════════════════════════════════════════════════
def build():
    doc = setup_doc()

    build_cover(doc)
    build_toc(doc)

    build_part1_intro(doc)
    build_section1(doc)
    build_section2(doc)
    build_section3(doc)
    build_section4(doc)

    build_part2_intro(doc)
    build_section5(doc)
    build_section6(doc)
    build_section7(doc)
    build_section8(doc)
    build_section9(doc)

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    build()
